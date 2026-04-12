# -*- coding: utf-8 -*-
"""
Document Parser Component

Extracts textual content from documents of various formats (PDF, DOCX, TXT, Markdown).
Provides parsed text and metadata for further processing.
"""

from datetime import datetime
from typing import Optional
import logging

from pydantic import BaseModel, Field
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

from ragindexer.FileScanner import FileInfo, FileFormat

logger = logging.getLogger(__name__)


class DocumentMetadata(BaseModel):
    """
    Metadata extracted from a document.

    Attributes:
        title: Document title (if available)
        author: Document author (if available)
        page_count: Number of pages (for PDF/DOCX)
        document: Original file path
        format: Document format (pdf, docx, txt, md)
        extraction_time: When the document was parsed
    """

    title: Optional[str] = None
    author: Optional[str] = None
    page_count: Optional[int] = None
    document: str
    format: FileFormat
    extraction_time: datetime = Field(default_factory=datetime.now)


class ParsedDocument(BaseModel):
    """
    Result of parsing a document.

    Attributes:
        content: Extracted text content
        metadata: Document metadata
        file_info: Original FileInfo from scanner
        character_count: Total character count of extracted content
    """

    content: str
    metadata: DocumentMetadata
    file_info: FileInfo
    character_count: int = Field(default=0)

    def __init__(self, **data):
        super().__init__(**data)
        self.character_count = len(self.content)


class DocumentParser:
    """
    Parses documents in various formats and extracts text content.

    Supports: PDF, DOCX, DOC (as DOCX), TXT, Markdown

    For PDF files, extracts text directly. OCR support can be added for scanned PDFs.
    """

    def __init__(self, logger_instance: Optional[logging.Logger] = None):
        """
        Initialize the DocumentParser.

        Args:
            logger_instance: Logger to use (defaults to module logger)
        """
        self.logger = logger_instance or logger

    def parse(self, file_info: FileInfo) -> ParsedDocument:
        """
        Parse a document and extract its text content.

        Args:
            file_info: FileInfo object from FileScanner

        Returns:
            ParsedDocument with extracted content and metadata

        Raises:
            ValueError: If file format is not supported
            IOError: If file cannot be read
        """
        if not file_info.absolute_path.exists():
            raise IOError(f"File does not exist: {file_info.absolute_path}")

        self.logger.info(f"Parsing document: {file_info.relative_path}")

        try:
            if file_info.format == FileFormat.PDF:
                content, metadata = self._parse_pdf(file_info)
            elif file_info.format == FileFormat.DOCX:
                content, metadata = self._parse_docx(file_info)
            elif file_info.format == FileFormat.DOC:
                content, metadata = self._parse_docx(file_info)
            elif file_info.format == FileFormat.TXT:
                content, metadata = self._parse_txt(file_info)
            elif file_info.format == FileFormat.MARKDOWN:
                content, metadata = self._parse_markdown(file_info)
            else:
                raise ValueError(f"Unsupported format: {file_info.format}")

            parsed_doc = ParsedDocument(
                content=content,
                metadata=metadata,
                file_info=file_info,
            )

            self.logger.info(
                f"Successfully parsed {file_info.relative_path}: "
                f"{parsed_doc.character_count} characters"
            )

            return parsed_doc

        except Exception as e:
            self.logger.error(f"Failed to parse {file_info.relative_path}: {e}")
            raise

    def _parse_pdf(self, file_info: FileInfo) -> tuple[str, DocumentMetadata]:
        """
        Extract text from a PDF file.

        Args:
            file_info: FileInfo object

        Returns:
            Tuple of (text content, DocumentMetadata)
        """
        content_parts = []
        page_count = 0

        try:
            with open(file_info.absolute_path, "rb") as f:
                reader = PdfReader(f)
                page_count = len(reader.pages)

                # Extract metadata from PDF
                metadata_dict = reader.metadata or {}
                title = metadata_dict.get("/Title")
                author = metadata_dict.get("/Author")

                # Extract text from each page
                for page_num, page in enumerate(reader.pages):
                    try:
                        text = page.extract_text()
                        if text:
                            content_parts.append(text)
                    except Exception as e:
                        self.logger.warning(
                            f"Failed to extract text from page {page_num + 1} "
                            f"of {file_info.relative_path}: {e}"
                        )

        except Exception as e:
            self.logger.error(f"Failed to read PDF {file_info.relative_path}: {e}")
            raise IOError(f"Cannot read PDF file: {e}")

        content = "\n".join(content_parts)

        metadata = DocumentMetadata(
            title=title,
            author=author,
            page_count=page_count,
            document=str(file_info.relative_path),
            format=file_info.format,
        )

        return content, metadata

    def _parse_docx(self, file_info: FileInfo) -> tuple[str, DocumentMetadata]:
        """
        Extract text from a DOCX file.

        Args:
            file_info: FileInfo object

        Returns:
            Tuple of (text content, DocumentMetadata)
        """
        content_parts = []

        try:
            doc = DocxDocument(str(file_info.absolute_path))

            # Extract title from core properties
            core_props = doc.core_properties
            title = core_props.title
            author = core_props.author

            # Extract text from all paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    if any(row_text):
                        content_parts.append(" | ".join(row_text))

        except Exception as e:
            self.logger.error(f"Failed to read DOCX {file_info.relative_path}: {e}")
            raise IOError(f"Cannot read DOCX file: {e}")

        content = "\n".join(content_parts)

        metadata = DocumentMetadata(
            title=title,
            author=author,
            page_count=None,  # DOCX doesn't have clear page count
            document=str(file_info.relative_path),
            format=file_info.format,
        )

        return content, metadata

    def _parse_txt(self, file_info: FileInfo) -> tuple[str, DocumentMetadata]:
        """
        Extract text from a TXT file.

        Args:
            file_info: FileInfo object

        Returns:
            Tuple of (text content, DocumentMetadata)
        """
        try:
            with open(file_info.absolute_path, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            try:
                with open(file_info.absolute_path, "r", encoding="latin-1") as f:
                    content = f.read()
            except Exception as e:
                self.logger.error(f"Failed to read TXT {file_info.relative_path}: {e}")
                raise IOError(f"Cannot read TXT file: {e}")
        except Exception as e:
            self.logger.error(f"Failed to read TXT {file_info.relative_path}: {e}")
            raise IOError(f"Cannot read TXT file: {e}")

        metadata = DocumentMetadata(
            title=None,
            author=None,
            page_count=None,
            document=str(file_info.relative_path),
            format=file_info.format,
        )

        return content, metadata

    def _parse_markdown(self, file_info: FileInfo) -> tuple[str, DocumentMetadata]:
        """
        Extract text from a Markdown file.

        Args:
            file_info: FileInfo object

        Returns:
            Tuple of (text content, DocumentMetadata)
        """
        try:
            with open(file_info.absolute_path, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            try:
                with open(file_info.absolute_path, "r", encoding="latin-1") as f:
                    content = f.read()
            except Exception as e:
                self.logger.error(f"Failed to read Markdown {file_info.relative_path}: {e}")
                raise IOError(f"Cannot read Markdown file: {e}")
        except Exception as e:
            self.logger.error(f"Failed to read Markdown {file_info.relative_path}: {e}")
            raise IOError(f"Cannot read Markdown file: {e}")

        metadata = DocumentMetadata(
            title=None,
            author=None,
            page_count=None,
            document=str(file_info.relative_path),
            format=file_info.format,
        )

        return content, metadata
