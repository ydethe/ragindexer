from typing import Any, Dict, Iterable, Tuple

import docx

from .. import logger
from .ADocument import ADocument


class DocDocument(ADocument):
    def iterate_raw_text(self) -> Iterable[Tuple[int, str, Dict[str, Any]]]:
        try:
            doc = docx.Document(str(self.get_abs_path()))
        except Exception:
            logger.warning("Error while reading the file. Skipping")
            return None, {"ocr_used": False}

        page_count = sum(p.contains_page_break for p in doc.paragraphs) + 1
        logger.info(f"Reading {page_count} pages doc file")
        avct = -1
        for k_page, p in enumerate(doc.paragraphs):
            new_avct = int(k_page / page_count * 100 / 10)
            if new_avct != avct:
                logger.info(f"Lecture page {k_page+1}/{page_count}")
                avct = new_avct

            yield k_page, "\n".join(p.text).strip(), {"ocr_used": False}
