# -*- coding: utf-8 -*-
"""
Unit tests for DocumentParser component.

Tests parsing of various document formats:
- PDF
- DOCX
- TXT
- Markdown
"""

import pytest
from pathlib import Path
from datetime import datetime
from tempfile import TemporaryDirectory

from ragindexer.DocumentParser import (
    DocumentParser,
    ParsedDocument,
)
from ragindexer.FileScanner import FileInfo, FileFormat


@pytest.fixture
def temp_docs_dir():
    """Create temporary directory with test documents."""
    with TemporaryDirectory() as tmpdir:
        yield Path(tmpdir)


@pytest.fixture
def document_parser():
    """Create DocumentParser instance."""
    return DocumentParser()


@pytest.fixture
def sample_txt_file(temp_docs_dir):
    """Create a sample TXT file."""
    txt_file = temp_docs_dir / "sample.txt"
    content = "Hello World\nThis is a test document.\nWith multiple lines."
    txt_file.write_text(content, encoding="utf-8")
    return txt_file


@pytest.fixture
def sample_markdown_file(temp_docs_dir):
    """Create a sample Markdown file."""
    md_file = temp_docs_dir / "sample.md"
    content = "# Title\n\nThis is a markdown document.\n\n## Section\n\nWith content."
    md_file.write_text(content, encoding="utf-8")
    return md_file


@pytest.fixture
def sample_docx_file(temp_docs_dir):
    """Create a sample DOCX file."""
    try:
        from docx import Document as DocxDocument

        docx_file = temp_docs_dir / "sample.docx"
        doc = DocxDocument()

        # Add title
        doc.core_properties.title = "Test Document"
        doc.core_properties.author = "Test Author"

        # Add paragraphs
        doc.add_paragraph("This is a test DOCX document.")
        doc.add_paragraph("With multiple paragraphs.")

        # Add table
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Header 1"
        table.rows[0].cells[1].text = "Header 2"
        table.rows[1].cells[0].text = "Data 1"
        table.rows[1].cells[1].text = "Data 2"

        doc.save(str(docx_file))
        return docx_file
    except ImportError:
        pytest.skip("python-docx not available")


@pytest.fixture
def sample_pdf_file():
    """Skip PDF file test - requires reportlab."""
    pytest.skip("PDF test requires reportlab dependency")


class TestDocumentParserTXT:
    """Test TXT file parsing."""

    def test_parse_txt_file(self, document_parser, sample_txt_file):
        """Test parsing a TXT file."""
        file_info = FileInfo(
            relative_path="sample.txt",
            absolute_path=sample_txt_file,
            format=FileFormat.TXT,
            file_size=sample_txt_file.stat().st_size,
            modified_time=datetime.now(),
            file_hash="testhash",
        )

        result = document_parser.parse(file_info)

        assert isinstance(result, ParsedDocument)
        assert "Hello World" in result.content
        assert "test document" in result.content
        assert result.character_count > 0
        assert result.metadata.format == FileFormat.TXT
        assert result.metadata.document == "sample.txt"

    def test_parse_txt_preserves_content(self, document_parser, sample_txt_file):
        """Test that TXT parsing preserves all content."""
        file_info = FileInfo(
            relative_path="sample.txt",
            absolute_path=sample_txt_file,
            format=FileFormat.TXT,
            file_size=sample_txt_file.stat().st_size,
            modified_time=datetime.now(),
            file_hash="testhash",
        )

        result = document_parser.parse(file_info)
        original_content = sample_txt_file.read_text(encoding="utf-8")

        assert result.content == original_content


class TestDocumentParserMarkdown:
    """Test Markdown file parsing."""

    def test_parse_markdown_file(self, document_parser, sample_markdown_file):
        """Test parsing a Markdown file."""
        file_info = FileInfo(
            relative_path="sample.md",
            absolute_path=sample_markdown_file,
            format=FileFormat.MARKDOWN,
            file_size=sample_markdown_file.stat().st_size,
            modified_time=datetime.now(),
            file_hash="testhash",
        )

        result = document_parser.parse(file_info)

        assert isinstance(result, ParsedDocument)
        assert "Title" in result.content
        assert "markdown" in result.content
        assert result.character_count > 0
        assert result.metadata.format == FileFormat.MARKDOWN

    def test_parse_markdown_preserves_structure(self, document_parser, sample_markdown_file):
        """Test that Markdown parsing preserves structure."""
        file_info = FileInfo(
            relative_path="sample.md",
            absolute_path=sample_markdown_file,
            format=FileFormat.MARKDOWN,
            file_size=sample_markdown_file.stat().st_size,
            modified_time=datetime.now(),
            file_hash="testhash",
        )

        result = document_parser.parse(file_info)
        original_content = sample_markdown_file.read_text(encoding="utf-8")

        assert result.content == original_content
        assert "#" in result.content  # Markdown syntax preserved


class TestDocumentParserDOCX:
    """Test DOCX file parsing."""

    def test_parse_docx_file(self, document_parser, sample_docx_file):
        """Test parsing a DOCX file."""
        file_info = FileInfo(
            relative_path="sample.docx",
            absolute_path=sample_docx_file,
            format=FileFormat.DOCX,
            file_size=sample_docx_file.stat().st_size,
            modified_time=datetime.now(),
            file_hash="testhash",
        )

        result = document_parser.parse(file_info)

        assert isinstance(result, ParsedDocument)
        assert "test DOCX document" in result.content
        assert "paragraphs" in result.content
        assert result.character_count > 0
        assert result.metadata.title == "Test Document"
        assert result.metadata.author == "Test Author"
        assert result.metadata.format == FileFormat.DOCX

    def test_parse_docx_extracts_tables(self, document_parser, sample_docx_file):
        """Test that DOCX parsing extracts table content."""
        file_info = FileInfo(
            relative_path="sample.docx",
            absolute_path=sample_docx_file,
            format=FileFormat.DOCX,
            file_size=sample_docx_file.stat().st_size,
            modified_time=datetime.now(),
            file_hash="testhash",
        )

        result = document_parser.parse(file_info)

        assert "Header 1" in result.content
        assert "Header 2" in result.content
        assert "Data 1" in result.content
        assert "Data 2" in result.content


class TestDocumentParserErrors:
    """Test error handling."""

    def test_parse_nonexistent_file(self, document_parser):
        """Test parsing a non-existent file raises IOError."""
        file_info = FileInfo(
            relative_path="nonexistent.txt",
            absolute_path=Path("/nonexistent/file.txt"),
            format=FileFormat.TXT,
            file_size=0,
            modified_time=datetime.now(),
            file_hash="testhash",
        )

        with pytest.raises(IOError):
            document_parser.parse(file_info)

    def test_parse_invalid_encoding_txt(self, document_parser, temp_docs_dir):
        """Test parsing TXT file with different encoding."""
        txt_file = temp_docs_dir / "latin1.txt"
        # Write with latin-1 encoding
        txt_file.write_text("Héllo Wörld", encoding="latin-1")

        file_info = FileInfo(
            relative_path="latin1.txt",
            absolute_path=txt_file,
            format=FileFormat.TXT,
            file_size=txt_file.stat().st_size,
            modified_time=datetime.now(),
            file_hash="testhash",
        )

        result = document_parser.parse(file_info)
        assert result.character_count > 0


class TestDocumentParserMetadata:
    """Test metadata extraction."""

    def test_metadata_includes_file_info(self, document_parser, sample_txt_file):
        """Test that metadata includes reference to FileInfo."""
        file_info = FileInfo(
            relative_path="sample.txt",
            absolute_path=sample_txt_file,
            format=FileFormat.TXT,
            file_size=sample_txt_file.stat().st_size,
            modified_time=datetime.now(),
            file_hash="testhash",
        )

        result = document_parser.parse(file_info)

        assert result.file_info == file_info
        assert result.metadata.document == file_info.relative_path

    def test_character_count_calculation(self, document_parser, sample_txt_file):
        """Test character count is correctly calculated."""
        file_info = FileInfo(
            relative_path="sample.txt",
            absolute_path=sample_txt_file,
            format=FileFormat.TXT,
            file_size=sample_txt_file.stat().st_size,
            modified_time=datetime.now(),
            file_hash="testhash",
        )

        result = document_parser.parse(file_info)
        expected_count = len(sample_txt_file.read_text(encoding="utf-8"))

        assert result.character_count == expected_count


class TestDocumentParserIntegration:
    """Integration tests with FileScanner."""

    def test_parse_after_file_scan(self, document_parser, sample_txt_file):
        """Test parsing documents after FileScanner detects them."""
        from ragindexer.FileScanner import FileScanner

        # Simulate FileScanner output
        scanner = FileScanner(sample_txt_file.parent)
        scan_result = scanner.scan()

        assert len(scan_result.files) > 0

        # Parse each detected file
        for rel_path, file_info in scan_result.files.items():
            result = document_parser.parse(file_info)
            assert isinstance(result, ParsedDocument)
            assert result.character_count > 0
